"""
Multi-format document parsers for ReasonableRAG
Supports: PDF, DOCX, TXT, CSV, XLSX, Images
"""
import os
import logging
from typing import Dict, Any, List, Optional
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


class DocumentParser:
    """Universal document parser for multiple formats"""
    
    def __init__(self):
        self.parsers = {
            '.pdf': self._parse_pdf,
            '.txt': self._parse_txt,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,
            '.csv': self._parse_csv,
            '.xlsx': self._parse_excel,
            '.xls': self._parse_excel,
            '.png': self._parse_image,
            '.jpg': self._parse_image,
            '.jpeg': self._parse_image,
            '.gif': self._parse_image,
            '.bmp': self._parse_image,
            '.tiff': self._parse_image
        }
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract content
        
        Args:
            file_path: Path to document
        
        Returns:
            Dictionary containing:
                - text: Extracted text content
                - metadata: Document metadata
                - tables: Extracted tables (if any)
                - success: Boolean success flag
                - error: Error message (if failed)
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.parsers:
            return {
                'text': '',
                'metadata': {},
                'tables': [],
                'success': False,
                'error': f'Unsupported file format: {ext}'
            }
        
        try:
            result = self.parsers[ext](file_path)
            result['success'] = True
            result['file_path'] = file_path
            result['file_name'] = os.path.basename(file_path)
            result['file_type'] = ext
            logger.info(f"Successfully parsed {file_path}")
            return result
        
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            return {
                'text': '',
                'metadata': {},
                'tables': [],
                'success': False,
                'error': str(e),
                'file_path': file_path
            }
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF documents"""
        doc = fitz.open(file_path)
        text_content = []
        tables = []
        
        metadata = {
            'pages': len(doc),
            'author': doc.metadata.get('author', ''),
            'title': doc.metadata.get('title', ''),
            'subject': doc.metadata.get('subject', '')
        }
        
        for page_num, page in enumerate(doc):
            # Extract text
            text = page.get_text()
            text_content.append(f"--- Page {page_num + 1} ---\n{text}")
            
            # Try to extract tables
            try:
                page_tables = page.find_tables()
                if page_tables:
                    for table in page_tables:
                        tables.append({
                            'page': page_num + 1,
                            'data': table.extract()
                        })
            except:
                pass
        
        doc.close()
        
        return {
            'text': '\n\n'.join(text_content),
            'metadata': metadata,
            'tables': tables
        }
    
    def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """Parse text files"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                return {
                    'text': text,
                    'metadata': {'encoding': encoding},
                    'tables': []
                }
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode text file with any standard encoding")
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX/DOC documents"""
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract tables
        tables = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                'table_index': table_idx,
                'data': table_data
            })
        
        # Metadata
        core_props = doc.core_properties
        metadata = {
            'author': core_props.author or '',
            'title': core_props.title or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else ''
        }
        
        return {
            'text': '\n\n'.join(text_content),
            'metadata': metadata,
            'tables': tables
        }
    
    def _parse_csv(self, file_path: str) -> Dict[str, Any]:
        """Parse CSV files"""
        # Try different encodings and delimiters
        encodings = ['utf-8', 'latin-1', 'cp1252']
        delimiters = [',', ';', '\t', '|']
        
        df = None
        used_encoding = None
        used_delimiter = None
        
        for encoding in encodings:
            for delimiter in delimiters:
                try:
                    df = pd.read_csv(file_path, encoding=encoding, 
                                   delimiter=delimiter, on_bad_lines='skip')
                    if len(df.columns) > 1:  # Valid CSV should have multiple columns
                        used_encoding = encoding
                        used_delimiter = delimiter
                        break
                except:
                    continue
            if df is not None:
                break
        
        if df is None:
            raise ValueError("Could not parse CSV file")
        
        # Convert to text representation
        text_parts = [f"CSV Data with {len(df)} rows and {len(df.columns)} columns\n"]
        text_parts.append("Columns: " + ", ".join(df.columns.astype(str).tolist()))
        text_parts.append("\nData Preview:\n" + df.head(100).to_string())
        
        # Check for messy data indicators
        is_messy = self._detect_messy_csv(df)
        
        metadata = {
            'rows': len(df),
            'columns': len(df.columns),
            'column_names': df.columns.tolist(),
            'encoding': used_encoding,
            'delimiter': used_delimiter,
            'is_messy': is_messy
        }
        
        return {
            'text': '\n\n'.join(text_parts),
            'metadata': metadata,
            'tables': [{'data': df.values.tolist(), 'columns': df.columns.tolist()}]
        }
    
    def _parse_excel(self, file_path: str) -> Dict[str, Any]:
        """Parse Excel files (XLSX/XLS)"""
        excel_file = pd.ExcelFile(file_path)
        
        text_parts = []
        all_tables = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
            text_parts.append("Columns: " + ", ".join(df.columns.astype(str).tolist()))
            text_parts.append(df.head(50).to_string())
            
            all_tables.append({
                'sheet_name': sheet_name,
                'data': df.values.tolist(),
                'columns': df.columns.tolist()
            })
        
        # Check for messy data
        first_df = pd.read_excel(file_path, sheet_name=excel_file.sheet_names[0])
        is_messy = self._detect_messy_csv(first_df)
        
        metadata = {
            'sheets': excel_file.sheet_names,
            'sheet_count': len(excel_file.sheet_names),
            'is_messy': is_messy
        }
        
        return {
            'text': '\n\n'.join(text_parts),
            'metadata': metadata,
            'tables': all_tables
        }
    
    def _parse_image(self, file_path: str) -> Dict[str, Any]:
        """Parse images using OCR"""
        try:
            image = Image.open(file_path)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            metadata = {
                'size': image.size,
                'mode': image.mode,
                'format': image.format
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'tables': []
            }
        
        except Exception as e:
            logger.warning(f"OCR failed for {file_path}, returning empty text: {e}")
            return {
                'text': '',
                'metadata': {'error': 'OCR failed'},
                'tables': []
            }
    
    def _detect_messy_csv(self, df: pd.DataFrame) -> bool:
        """
        Detect if CSV/Excel data is messy or unorganized
        
        Indicators:
        - Missing column names (Unnamed columns)
        - Too many columns with mostly null values
        - Inconsistent data types in columns
        - Very few rows
        """
        messy_indicators = 0
        
        # Check for unnamed columns
        unnamed_cols = sum(1 for col in df.columns if 'Unnamed' in str(col))
        if unnamed_cols > len(df.columns) * 0.3:  # More than 30% unnamed
            messy_indicators += 1
        
        # Check for columns with mostly null values
        null_cols = sum(1 for col in df.columns if df[col].isna().sum() > len(df) * 0.5)
        if null_cols > len(df.columns) * 0.3:
            messy_indicators += 1
        
        # Check if too few rows
        if len(df) < 3:
            messy_indicators += 1
        
        # Check if too many columns
        if len(df.columns) > 50:
            messy_indicators += 1
        
        return messy_indicators >= 2


logger.info("Document parsers initialized")

